import csv
import io
from dataclasses import dataclass, field

import openpyxl
import xlrd
from docx import Document as WordDocument
from pypdf import PdfReader


class DocumentExtractionError(ValueError):
    """Raised when document text cannot be extracted safely."""


@dataclass(slots=True)
class ExtractedSection:
    text: str
    page_number: int | None = None
    section_title: str | None = None
    metadata: dict[str, str | int] = field(default_factory=dict)


@dataclass(slots=True)
class ExtractedDocument:
    sections: list[ExtractedSection]
    page_count: int | None = None

    @property
    def text(self) -> str:
        return "\n\n".join(
            section.text for section in self.sections if section.text.strip()
        )


def normalize_text(value: str) -> str:
    lines = [" ".join(line.split()) for line in value.replace("\x00", "").splitlines()]

    return "\n".join(line for line in lines if line).strip()


def decode_text_file(content: bytes) -> str:
    encodings = (
        "utf-8-sig",
        "utf-8",
        "utf-16",
        "cp1252",
    )

    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentExtractionError("The text file encoding is not supported.")


def extract_txt(content: bytes) -> ExtractedDocument:
    text = normalize_text(decode_text_file(content))

    if not text:
        raise DocumentExtractionError(
            "The text document contains no extractable content."
        )

    return ExtractedDocument(
        sections=[
            ExtractedSection(
                text=text,
                section_title="Text document",
            )
        ]
    )


def extract_csv(content: bytes) -> ExtractedDocument:
    decoded_content = decode_text_file(content)

    try:
        rows = list(csv.reader(io.StringIO(decoded_content)))
    except csv.Error as exc:
        raise DocumentExtractionError("The CSV document could not be parsed.") from exc

    normalized_rows = [
        " | ".join(normalize_text(str(cell)) for cell in row)
        for row in rows
        if any(str(cell).strip() for cell in row)
    ]

    if not normalized_rows:
        raise DocumentExtractionError(
            "The CSV document contains no extractable content."
        )

    return ExtractedDocument(
        sections=[
            ExtractedSection(
                text="\n".join(normalized_rows),
                section_title="CSV data",
                metadata={
                    "row_count": len(normalized_rows),
                },
            )
        ]
    )


def extract_pdf(content: bytes) -> ExtractedDocument:
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("The PDF document could not be opened.") from exc

    if reader.is_encrypted:
        try:
            decryption_result = reader.decrypt("")
        except Exception as exc:
            raise DocumentExtractionError(
                "Encrypted PDF documents are not supported."
            ) from exc

        if decryption_result == 0:
            raise DocumentExtractionError("Encrypted PDF documents are not supported.")

    sections: list[ExtractedSection] = []

    for page_index, page in enumerate(
        reader.pages,
        start=1,
    ):
        try:
            extracted_text = page.extract_text() or ""
        except Exception as exc:
            raise DocumentExtractionError(
                "Text could not be extracted from the PDF."
            ) from exc

        normalized_text = normalize_text(extracted_text)

        if normalized_text:
            sections.append(
                ExtractedSection(
                    text=normalized_text,
                    page_number=page_index,
                    section_title=f"Page {page_index}",
                )
            )

    if not sections:
        raise DocumentExtractionError(
            "The PDF contains no extractable text. Scanned PDFs require OCR."
        )

    return ExtractedDocument(
        sections=sections,
        page_count=len(reader.pages),
    )


def extract_docx(content: bytes) -> ExtractedDocument:
    try:
        document = WordDocument(io.BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("The Word document could not be opened.") from exc

    sections: list[ExtractedSection] = []
    current_title: str | None = None
    current_paragraphs: list[str] = []

    def save_section() -> None:
        if not current_paragraphs:
            return

        section_text = normalize_text("\n".join(current_paragraphs))

        if section_text:
            sections.append(
                ExtractedSection(
                    text=section_text,
                    section_title=current_title,
                )
            )

        current_paragraphs.clear()

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)

        if not text:
            continue

        style_name = (
            paragraph.style.name.lower()
            if paragraph.style and paragraph.style.name
            else ""
        )

        if style_name.startswith("heading"):
            save_section()
            current_title = text
        else:
            current_paragraphs.append(text)

    for table_index, table in enumerate(
        document.tables,
        start=1,
    ):
        table_rows = []

        for row in table.rows:
            values = [normalize_text(cell.text) for cell in row.cells]
            table_rows.append(" | ".join(values))

        table_text = normalize_text("\n".join(table_rows))

        if table_text:
            current_paragraphs.append(f"Table {table_index}\n{table_text}")

    save_section()

    if not sections:
        raise DocumentExtractionError(
            "The Word document contains no extractable content."
        )

    return ExtractedDocument(sections=sections)


def worksheet_to_text(
    worksheet: openpyxl.worksheet.worksheet.Worksheet,
) -> str:
    rows: list[str] = []

    for row in worksheet.iter_rows(values_only=True):
        values = [
            normalize_text(str(value)) if value is not None else "" for value in row
        ]

        if any(values):
            rows.append(" | ".join(values))

    return "\n".join(rows)


def extract_xlsx(content: bytes) -> ExtractedDocument:
    workbook: openpyxl.Workbook | None = None

    try:
        workbook = openpyxl.load_workbook(
            io.BytesIO(content),
            read_only=True,
            data_only=True,
        )

        sections = []

        for worksheet in workbook.worksheets:
            sheet_text = worksheet_to_text(worksheet)

            if sheet_text:
                sections.append(
                    ExtractedSection(
                        text=sheet_text,
                        section_title=worksheet.title,
                        metadata={
                            "sheet_name": worksheet.title,
                        },
                    )
                )
    except Exception as exc:
        raise DocumentExtractionError(
            "The Excel workbook could not be opened."
        ) from exc
    finally:
        if workbook is not None:
            workbook.close()

    if not sections:
        raise DocumentExtractionError(
            "The Excel workbook contains no extractable content."
        )

    return ExtractedDocument(sections=sections)


def extract_xls(content: bytes) -> ExtractedDocument:
    try:
        workbook = xlrd.open_workbook(file_contents=content)
    except Exception as exc:
        raise DocumentExtractionError(
            "The legacy Excel workbook could not be opened."
        ) from exc

    sections: list[ExtractedSection] = []

    for sheet in workbook.sheets():
        rows: list[str] = []

        for row_index in range(sheet.nrows):
            values = [
                normalize_text(
                    str(
                        sheet.cell_value(
                            row_index,
                            column_index,
                        )
                    )
                )
                for column_index in range(sheet.ncols)
            ]

            if any(values):
                rows.append(" | ".join(values))

        if rows:
            sections.append(
                ExtractedSection(
                    text="\n".join(rows),
                    section_title=sheet.name,
                    metadata={
                        "sheet_name": sheet.name,
                    },
                )
            )

    if not sections:
        raise DocumentExtractionError(
            "The legacy Excel workbook contains no extractable content."
        )

    return ExtractedDocument(sections=sections)


def extract_document_text(
    *,
    content: bytes,
    file_extension: str,
) -> ExtractedDocument:
    if not content:
        raise DocumentExtractionError("The document is empty.")

    normalized_extension = file_extension.lower().lstrip(".")

    extractors = {
        "txt": extract_txt,
        "csv": extract_csv,
        "pdf": extract_pdf,
        "docx": extract_docx,
        "xlsx": extract_xlsx,
        "xls": extract_xls,
    }

    extractor = extractors.get(normalized_extension)

    if extractor is None:
        raise DocumentExtractionError(
            "Text extraction is not supported for this file type."
        )

    return extractor(content)
