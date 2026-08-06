import io

import openpyxl
import pytest
from docx import Document as WordDocument

from services.document.text_extractor import (
    DocumentExtractionError,
    extract_document_text,
)


def test_extract_txt() -> None:
    result = extract_document_text(
        content=(b"First line.\n\nSecond line."),
        file_extension="txt",
    )

    assert "First line." in result.text
    assert "Second line." in result.text


def test_extract_utf8_bom_text() -> None:
    result = extract_document_text(
        content=("\ufeffCompany policy".encode("utf-8")),
        file_extension="txt",
    )

    assert result.text == "Company policy"


def test_extract_csv() -> None:
    result = extract_document_text(
        content=(b"name,status\nPolicy A,active\nPolicy B,draft\n"),
        file_extension="csv",
    )

    assert "name | status" in result.text
    assert "Policy A | active" in result.text


def test_extract_docx() -> None:
    stream = io.BytesIO()
    document = WordDocument()
    document.add_heading(
        "Employee Policy",
        level=1,
    )
    document.add_paragraph("Employees must follow the policy.")
    document.save(stream)

    result = extract_document_text(
        content=stream.getvalue(),
        file_extension="docx",
    )

    assert result.sections[0].section_title == ("Employee Policy")
    assert "Employees must follow" in result.text


def test_extract_xlsx() -> None:
    stream = io.BytesIO()
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Policies"
    worksheet.append(["name", "status"])
    worksheet.append(["Leave Policy", "active"])
    workbook.save(stream)
    workbook.close()

    result = extract_document_text(
        content=stream.getvalue(),
        file_extension="xlsx",
    )

    assert result.sections[0].section_title == ("Policies")
    assert "Leave Policy | active" in result.text


def test_empty_document_is_rejected() -> None:
    with pytest.raises(
        DocumentExtractionError,
        match="document is empty",
    ):
        extract_document_text(
            content=b"",
            file_extension="txt",
        )


def test_unsupported_extension_is_rejected() -> None:
    with pytest.raises(
        DocumentExtractionError,
        match="not supported",
    ):
        extract_document_text(
            content=b"content",
            file_extension="exe",
        )


def test_empty_text_is_rejected() -> None:
    with pytest.raises(
        DocumentExtractionError,
        match="no extractable content",
    ):
        extract_document_text(
            content=b"   \n\n  ",
            file_extension="txt",
        )
